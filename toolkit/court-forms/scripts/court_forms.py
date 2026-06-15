#!/usr/bin/env python3
"""Sync, search, and download public Korean court form templates.

The source is the public 대한민국 법원 전자소송포털 양식모음 page.
Only metadata is synced by default; template files are downloaded on demand.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import sqlite3
import subprocess
import sys
import time
import urllib.error
import urllib.parse
import urllib.request
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from xml.etree import ElementTree


ROOT = Path(os.path.expanduser("~/court-forms"))
DB_PATH = ROOT / "db" / "forms.db"
FILES_DIR = ROOT / "files"
SOURCE_URL = "https://ecfs.scourt.go.kr/psp/index.on?m=PSP720M24"
COPYRIGHT_POLICY_URL = "https://ecfs.scourt.go.kr/psp/index.on?m=PSP023M01"
LIST_API_URL = "https://ecfs.scourt.go.kr/psp/psp720/selectNboardList.on"
DOWNLOAD_URL = "https://file.scourt.go.kr/AttachDownload"
USER_AGENT = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/149.0.0.0 Safari/537.36"
)
TOKEN_RE = re.compile(r"[0-9A-Za-z가-힣_]+")
BAD_FILENAME_RE = re.compile(r'[\\/:*?"<>|\x00-\x1f]')
SPACE_RE = re.compile(r"[ \t]+")


class CourtFormsError(RuntimeError):
    """Expected operational failure."""


@dataclass(frozen=True)
class Attachment:
    kind: str
    original_name: str
    system_name: str
    download_url: str


def now_epoch_ms() -> int:
    return int(time.time() * 1000)


def sha1_short(value: str, length: int = 16) -> str:
    return hashlib.sha1(value.encode("utf-8")).hexdigest()[:length]


def connect(db_path: Path = DB_PATH) -> sqlite3.Connection:
    db_path.parent.mkdir(parents=True, exist_ok=True)
    con = sqlite3.connect(db_path)
    con.row_factory = sqlite3.Row
    con.execute("PRAGMA foreign_keys = ON")
    return con


def ensure_schema(con: sqlite3.Connection) -> None:
    con.executescript(
        """
        CREATE TABLE IF NOT EXISTS sync_runs (
          sync_id TEXT PRIMARY KEY,
          started_at TEXT DEFAULT CURRENT_TIMESTAMP,
          completed_at TEXT,
          source_url TEXT NOT NULL,
          total_count INTEGER,
          synced_forms INTEGER DEFAULT 0,
          downloaded_files INTEGER DEFAULT 0,
          status TEXT NOT NULL DEFAULT 'running',
          message TEXT
        );

        CREATE TABLE IF NOT EXISTS categories (
          category_code TEXT PRIMARY KEY,
          category_name TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS forms (
          form_id TEXT PRIMARY KEY,
          title TEXT NOT NULL,
          category_code TEXT,
          category_name TEXT,
          gubun INTEGER,
          source_url TEXT NOT NULL,
          source_page INTEGER,
          first_seen_at TEXT DEFAULT CURRENT_TIMESTAMP,
          last_seen_at TEXT DEFAULT CURRENT_TIMESTAMP
        );

        CREATE TABLE IF NOT EXISTS attachments (
          attachment_id TEXT PRIMARY KEY,
          form_id TEXT NOT NULL REFERENCES forms(form_id) ON DELETE CASCADE,
          kind TEXT NOT NULL,
          original_name TEXT NOT NULL,
          system_name TEXT NOT NULL,
          download_url TEXT NOT NULL,
          local_path TEXT,
          size_bytes INTEGER,
          sha256 TEXT,
          downloaded_at TEXT,
          UNIQUE(form_id, system_name)
        );

        CREATE VIRTUAL TABLE IF NOT EXISTS forms_fts USING fts5(
          form_id UNINDEXED,
          title,
          category_name,
          attachment_names,
          tokenize='unicode61'
        );
        """
    )
    con.commit()


def quote_cp949(value: str) -> str:
    return urllib.parse.quote(value, safe="", encoding="cp949", errors="replace")


def build_download_url(system_name: str, original_name: str) -> str:
    return (
        f"{DOWNLOAD_URL}?path=004"
        f"&file={urllib.parse.quote(system_name, safe='')}"
        f"&downFile={quote_cp949(original_name)}"
    )


def file_kind(filename: str) -> str:
    suffix = Path(filename).suffix.lower().lstrip(".")
    if suffix in {"doc", "docx"}:
        return suffix
    if suffix in {"hwp", "hwpx", "pdf"}:
        return suffix
    return suffix or "file"


def extract_attachments(row: dict[str, Any]) -> list[Attachment]:
    attachments: list[Attachment] = []
    for suffix in ("", "2", "3"):
        original = (row.get(f"fileOrgName{suffix}") or "").strip()
        system = (row.get(f"fileSysName{suffix}") or "").strip()
        if not original or not system:
            continue
        attachments.append(
            Attachment(
                kind=file_kind(original or system),
                original_name=original,
                system_name=system,
                download_url=build_download_url(system, original),
            )
        )
    return attachments


def form_id_for(row: dict[str, Any], attachments: list[Attachment]) -> str:
    material = "\0".join(
        [
            str(row.get("minGubun") or ""),
            str(row.get("title") or ""),
            "|".join(a.system_name for a in attachments),
        ]
    )
    return sha1_short(material, 18)


def attachment_id_for(form_id: str, attachment: Attachment) -> str:
    return sha1_short(f"{form_id}\0{attachment.system_name}\0{attachment.original_name}", 20)


def post_json(url: str, payload: dict[str, Any], timeout: int = 20) -> dict[str, Any]:
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    request = urllib.request.Request(
        url,
        data=body,
        method="POST",
        headers={
            "User-Agent": USER_AGENT,
            "Accept": "application/json",
            "Content-Type": "application/json;charset=UTF-8",
            "Origin": "https://ecfs.scourt.go.kr",
            "Referer": SOURCE_URL,
            "submissionid": "mf_pfwork_sbm_search",
            "sc-userid": "anonymous",
        },
    )
    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            raw = response.read()
    except urllib.error.URLError as exc:
        raise CourtFormsError(f"court forms list request failed: {exc}") from exc
    try:
        return json.loads(raw.decode("utf-8"))
    except json.JSONDecodeError as exc:
        sample = raw[:200].decode("utf-8", errors="replace")
        raise CourtFormsError(f"court forms list did not return JSON: {sample}") from exc


def list_payload(
    page_no: int,
    page_size: int,
    category_code: str = "",
    search_word: str = "",
    first: bool = False,
) -> dict[str, Any]:
    return {
        "dma_search": {
            "pageNo": page_no,
            "pageSize": page_size,
            "bfPageNo": "",
            "startRowNo": ((page_no - 1) * page_size) + 1,
            "totalCnt": "",
            "totalYn": "Y",
            "minGubun": category_code,
            "minGubun2": "",
            "sName": "",
            "eName": "",
            "searchWord": search_word,
            "searchWord2": "",
            "firstYn": "Y" if first else "N",
        }
    }


def fetch_page(
    page_no: int,
    page_size: int,
    category_code: str = "",
    search_word: str = "",
) -> dict[str, Any]:
    payload = list_payload(page_no, page_size, category_code, search_word, first=page_no == 1)
    result = post_json(LIST_API_URL, payload)
    if result.get("status") != 200:
        raise CourtFormsError(f"court forms API error: {result.get('message') or result}")
    return result.get("data") or {}


def upsert_categories(con: sqlite3.Connection, categories: list[dict[str, Any]]) -> None:
    for category in categories:
        code = (category.get("scode") or "").strip()
        name = (category.get("scodeKname") or "").strip()
        if not code or not name:
            continue
        con.execute(
            """
            INSERT INTO categories(category_code, category_name)
            VALUES (?, ?)
            ON CONFLICT(category_code) DO UPDATE SET
              category_name = excluded.category_name
            """,
            (code, name),
        )


def store_rows(
    con: sqlite3.Connection,
    rows: list[dict[str, Any]],
    category_names: dict[str, str],
    source_page: int,
) -> int:
    count = 0
    for row in rows:
        title = (row.get("title") or "").strip()
        if not title:
            continue
        attachments = extract_attachments(row)
        form_id = form_id_for(row, attachments)
        category_code = (row.get("minGubun") or "").strip()
        category_name = category_names.get(category_code, "")
        con.execute(
            """
            INSERT INTO forms(
              form_id, title, category_code, category_name, gubun, source_url, source_page
            )
            VALUES (?, ?, ?, ?, ?, ?, ?)
            ON CONFLICT(form_id) DO UPDATE SET
              title = excluded.title,
              category_code = excluded.category_code,
              category_name = excluded.category_name,
              gubun = excluded.gubun,
              source_url = excluded.source_url,
              source_page = excluded.source_page,
              last_seen_at = CURRENT_TIMESTAMP
            """,
            (
                form_id,
                title,
                category_code,
                category_name,
                row.get("gubun"),
                SOURCE_URL,
                source_page,
            ),
        )
        for attachment in attachments:
            con.execute(
                """
                INSERT INTO attachments(
                  attachment_id, form_id, kind, original_name, system_name, download_url
                )
                VALUES (?, ?, ?, ?, ?, ?)
                ON CONFLICT(attachment_id) DO UPDATE SET
                  form_id = excluded.form_id,
                  kind = excluded.kind,
                  original_name = excluded.original_name,
                  system_name = excluded.system_name,
                  download_url = excluded.download_url
                """,
                (
                    attachment_id_for(form_id, attachment),
                    form_id,
                    attachment.kind,
                    attachment.original_name,
                    attachment.system_name,
                    attachment.download_url,
                ),
            )
        count += 1
    return count


def rebuild_fts(con: sqlite3.Connection) -> None:
    con.execute("DELETE FROM forms_fts")
    con.execute(
        """
        INSERT INTO forms_fts(form_id, title, category_name, attachment_names)
        SELECT
          f.form_id,
          f.title,
          COALESCE(f.category_name, ''),
          COALESCE(GROUP_CONCAT(a.original_name, ' '), '')
        FROM forms f
        LEFT JOIN attachments a ON a.form_id = f.form_id
        GROUP BY f.form_id
        """
    )


def safe_filename(name: str) -> str:
    safe = BAD_FILENAME_RE.sub("_", name).strip(" .")
    return safe or "court-form"


def safe_dirname(value: str) -> str:
    safe = BAD_FILENAME_RE.sub("_", value)
    safe = SPACE_RE.sub("_", safe.strip())
    return safe[:120].strip("._") or "court-form"


def markdown_escape(value: str) -> str:
    return value.replace("\\", "\\\\").replace("|", "\\|")


def rel_link(target: Path, base_dir: Path) -> str:
    try:
        return urllib.parse.quote(str(target.relative_to(base_dir)).replace(os.sep, "/"))
    except ValueError:
        return urllib.parse.quote(str(target))


def download_bytes(url: str, timeout: int = 60) -> bytes:
    request = urllib.request.Request(
        url,
        headers={
            "User-Agent": USER_AGENT,
            "Referer": SOURCE_URL,
        },
    )
    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            data = response.read()
    except urllib.error.URLError as exc:
        raise CourtFormsError(f"download failed: {exc}") from exc
    if data.lstrip().startswith(b"<") and b"Web firewall" in data[:1000]:
        raise CourtFormsError("download was blocked by the source site's web firewall")
    return data


def download_attachment(
    con: sqlite3.Connection,
    attachment: sqlite3.Row,
    out_dir: Path,
    force: bool = False,
) -> dict[str, Any]:
    out_dir.mkdir(parents=True, exist_ok=True)
    target = out_dir / f"{attachment['attachment_id']}_{safe_filename(attachment['original_name'])}"
    if target.exists() and not force:
        data = target.read_bytes()
    else:
        data = download_bytes(attachment["download_url"])
        target.write_bytes(data)
    digest = hashlib.sha256(data).hexdigest()
    con.execute(
        """
        UPDATE attachments
        SET local_path = ?, size_bytes = ?, sha256 = ?, downloaded_at = CURRENT_TIMESTAMP
        WHERE attachment_id = ?
        """,
        (str(target), len(data), digest, attachment["attachment_id"]),
    )
    return {
        "attachment_id": attachment["attachment_id"],
        "kind": attachment["kind"],
        "original_name": attachment["original_name"],
        "path": str(target),
        "size_bytes": len(data),
        "sha256": digest,
    }


def run_text_command(command: list[str], timeout: int = 30) -> str:
    completed = subprocess.run(
        command,
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        timeout=timeout,
    )
    raw = completed.stdout
    for encoding in ("utf-8", "cp949", "euc-kr"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def extract_hwp_text(path: Path) -> tuple[str, str | None]:
    hwp5txt = shutil.which("hwp5txt")
    if hwp5txt:
        try:
            text = run_text_command([hwp5txt, str(path)])
            return text.strip(), None
        except Exception as exc:
            return "", f"hwp5txt failed: {exc}"
    hwpjs = shutil.which("hwpjs")
    if hwpjs:
        try:
            with tempfile_path(path.with_suffix(".tmp.md")) as tmp:
                subprocess.run(
                    [hwpjs, "to-markdown", str(path), "-o", str(tmp)],
                    check=True,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    timeout=60,
                )
                return tmp.read_text(encoding="utf-8", errors="replace").strip(), None
        except Exception as exc:
            return "", f"hwpjs failed: {exc}"
    return "", "no HWP text extractor found (install pyhwp/hwp5txt or @ohah/hwpjs)"


class tempfile_path:
    def __init__(self, path: Path):
        self.path = path

    def __enter__(self) -> Path:
        return self.path

    def __exit__(self, exc_type, exc, tb) -> None:
        try:
            self.path.unlink()
        except FileNotFoundError:
            pass


def extract_hwpx_text(path: Path) -> tuple[str, str | None]:
    chunks: list[str] = []
    try:
        with zipfile.ZipFile(path) as zf:
            names = sorted(
                n
                for n in zf.namelist()
                if n.lower().endswith(".xml") and ("section" in n.lower() or "contents/" in n.lower())
            )
            for name in names:
                try:
                    root = ElementTree.fromstring(zf.read(name))
                except ElementTree.ParseError:
                    continue
                for elem in root.iter():
                    if elem.text and elem.text.strip():
                        chunks.append(elem.text.strip())
    except Exception as exc:
        return "", f"hwpx parse failed: {exc}"
    return "\n".join(chunks).strip(), None if chunks else "no text found in hwpx"


def extract_docx_text(path: Path) -> tuple[str, str | None]:
    try:
        from docx import Document
    except Exception as exc:
        return "", f"python-docx unavailable: {exc}"
    try:
        doc = Document(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs).strip(), None
    except Exception as exc:
        return "", f"docx parse failed: {exc}"


def extract_pdf_text(path: Path) -> tuple[str, str | None]:
    try:
        import pdfplumber
    except Exception:
        pdfplumber = None
    if pdfplumber is not None:
        try:
            parts = []
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    if text.strip():
                        parts.append(text.strip())
            return "\n\n".join(parts).strip(), None
        except Exception as exc:
            return "", f"pdfplumber failed: {exc}"
    try:
        from pypdf import PdfReader
    except Exception as exc:
        return "", f"no PDF extractor found: {exc}"
    try:
        reader = PdfReader(str(path))
        return "\n\n".join((page.extract_text() or "").strip() for page in reader.pages).strip(), None
    except Exception as exc:
        return "", f"pypdf failed: {exc}"


def extract_text(path: Path, kind: str) -> tuple[str, str | None]:
    kind = kind.lower()
    if kind == "hwp":
        return extract_hwp_text(path)
    if kind == "hwpx":
        return extract_hwpx_text(path)
    if kind == "docx":
        return extract_docx_text(path)
    if kind == "pdf":
        return extract_pdf_text(path)
    if kind == "doc":
        return "", "legacy .doc extraction is not supported without LibreOffice/pandoc"
    return "", f"unsupported attachment kind: {kind}"


def preferred_text_attachment(attachments: list[dict[str, Any]]) -> dict[str, Any] | None:
    priority = {"hwp": 1, "hwpx": 2, "docx": 3, "pdf": 4, "doc": 5}
    available = [a for a in attachments if a.get("local_path")]
    available.sort(key=lambda a: priority.get(str(a.get("kind") or ""), 99))
    return available[0] if available else None


def form_rows(con: sqlite3.Connection) -> list[dict[str, Any]]:
    rows = con.execute(
        """
        SELECT *
        FROM forms
        ORDER BY category_name, title, form_id
        """
    ).fetchall()
    return [dict(r) for r in rows]


def copy_attachment_file(attachment: dict[str, Any], form_dir: Path) -> dict[str, Any]:
    local_path = attachment.get("local_path")
    if not local_path:
        return attachment
    src = Path(local_path)
    if not src.exists():
        return attachment
    original_dir = form_dir / "original"
    original_dir.mkdir(parents=True, exist_ok=True)
    target = original_dir / f"{attachment['attachment_id']}_{safe_filename(attachment['original_name'])}"
    if src.resolve() != target.resolve():
        shutil.copy2(src, target)
    updated = dict(attachment)
    updated["export_path"] = str(target)
    return updated


def write_form_markdown(
    form: dict[str, Any],
    attachments: list[dict[str, Any]],
    form_dir: Path,
    export_root: Path,
) -> tuple[Path, str | None]:
    form_dir.mkdir(parents=True, exist_ok=True)
    preferred = preferred_text_attachment(attachments)
    extracted_text = ""
    extraction_warning = None
    if preferred:
        source_path = Path(preferred.get("export_path") or preferred["local_path"])
        extracted_text, extraction_warning = extract_text(source_path, preferred["kind"])
    else:
        extraction_warning = "no local attachment file available"

    lines = [
        "---",
        f'form_id: "{form["form_id"]}"',
        f'title: "{str(form["title"]).replace(chr(34), chr(92) + chr(34))}"',
        f'category_code: "{form.get("category_code") or ""}"',
        f'category_name: "{form.get("category_name") or ""}"',
        f'source_url: "{form.get("source_url") or SOURCE_URL}"',
        f'copyright_policy_url: "{COPYRIGHT_POLICY_URL}"',
        'source_notice: "출처: 대한민국 법원 전자소송포털 양식모음"',
        "---",
        "",
        f"# {form['title']}",
        "",
        "> 출처: 대한민국 법원 전자소송포털 양식모음",
        f"> 원문: {form.get('source_url') or SOURCE_URL}",
        f"> 저작권보호정책: {COPYRIGHT_POLICY_URL}",
        "",
        "## 첨부파일",
        "",
        "| 형식 | 원본 파일명 | 로컬 경로 | SHA-256 |",
        "|---|---|---|---|",
    ]
    for attachment in attachments:
        path = attachment.get("export_path") or attachment.get("local_path") or ""
        link = f"[{Path(path).name}]({rel_link(Path(path), form_dir)})" if path else ""
        lines.append(
            "| "
            + " | ".join(
                [
                    markdown_escape(str(attachment.get("kind") or "")),
                    markdown_escape(str(attachment.get("original_name") or "")),
                    link,
                    markdown_escape(str(attachment.get("sha256") or "")),
                ]
            )
            + " |"
        )
    lines.extend(
        [
            "",
            "## 작성 활용 메모",
            "",
            "- 이 Markdown은 검색·초안 작성용 파생물입니다.",
            "- 제출·편집 기준은 위 공식 원본 HWP/PDF/DOC 파일입니다.",
            "- 사건번호, 법원, 당사자, 대리인, 주소, 신청취지 등은 사건별로 확인해 채우세요.",
            "",
            "## 추출 텍스트",
            "",
        ]
    )
    if extraction_warning:
        lines.extend([f"> 추출 경고: {extraction_warning}", ""])
    if extracted_text:
        lines.append(strip_trailing_whitespace(extracted_text))
    else:
        lines.append("_본문 텍스트를 자동 추출하지 못했습니다. 공식 원본 파일을 확인하세요._")
    lines.append("")

    target = form_dir / "index.md"
    target.write_text("\n".join(lines), encoding="utf-8")
    return target, extraction_warning


def strip_trailing_whitespace(text: str) -> str:
    return "\n".join(line.rstrip() for line in text.splitlines())


def write_export_readme(export_root: Path, form_count: int, attachment_count: int) -> None:
    readme = export_root / "README.md"
    readme.write_text(
        "\n".join(
            [
                "# 대한민국 법원 전자소송포털 양식모음 로컬 미러",
                "",
                f"- 양식 수: {form_count}",
                f"- 첨부파일 수: {attachment_count}",
                f"- 출처: {SOURCE_URL}",
                f"- 저작권보호정책: {COPYRIGHT_POLICY_URL}",
                "",
                "이 디렉터리는 공식 양식을 검색하고 초안 작성에 활용하기 위한 로컬 파생물입니다.",
                "Markdown은 검색·검토용이며, 제출 기준은 각 양식 폴더의 `original/` 원본 파일입니다.",
                "",
            ]
        ),
        encoding="utf-8",
    )


def manifest_attachment(attachment: dict[str, Any], export_root: Path) -> dict[str, Any]:
    export_path = attachment.get("export_path")
    return {
        "attachment_id": attachment.get("attachment_id"),
        "kind": attachment.get("kind"),
        "original_name": attachment.get("original_name"),
        "system_name": attachment.get("system_name"),
        "download_url": attachment.get("download_url"),
        "export_path": str(Path(export_path).relative_to(export_root)) if export_path else None,
        "size_bytes": attachment.get("size_bytes"),
        "sha256": attachment.get("sha256"),
        "downloaded_at": attachment.get("downloaded_at"),
    }


def export_markdown(args: argparse.Namespace) -> dict[str, Any]:
    con = connect(args.db)
    ensure_schema(con)
    export_root = args.output
    export_root.mkdir(parents=True, exist_ok=True)
    manifest_path = export_root / "forms.jsonl"
    exported = 0
    warnings = 0
    downloaded = 0
    attachment_count = 0
    download_errors: list[dict[str, str]] = []
    try:
        with manifest_path.open("w", encoding="utf-8") as manifest:
            forms = form_rows(con)
            for form in forms:
                attachments = attachments_for_form(con, form["form_id"])
                if args.download_missing:
                    refreshed = []
                    for attachment in attachments:
                        if not attachment.get("local_path"):
                            row = con.execute(
                                "SELECT * FROM attachments WHERE attachment_id = ?",
                                (attachment["attachment_id"],),
                            ).fetchone()
                            if row:
                                try:
                                    download_attachment(con, row, args.files_dir, force=args.force)
                                    downloaded += 1
                                except Exception as exc:
                                    if not args.continue_on_error:
                                        raise
                                    download_errors.append(
                                        {
                                            "attachment_id": row["attachment_id"],
                                            "original_name": row["original_name"],
                                            "error": str(exc),
                                        }
                                    )
                                if args.delay:
                                    time.sleep(args.delay)
                        refreshed.append(attachment["attachment_id"])
                    con.commit()
                    attachments = attachments_for_form(con, form["form_id"])
                category = safe_dirname(form.get("category_name") or form.get("category_code") or "uncategorized")
                title_slug = safe_dirname(form["title"])
                form_dir = export_root / "forms" / category / f"{form['form_id']}_{title_slug}"
                exported_attachments = [
                    copy_attachment_file(a, form_dir) if args.copy_files else dict(a)
                    for a in attachments
                ]
                md_path, warning = write_form_markdown(form, exported_attachments, form_dir, export_root)
                if warning:
                    warnings += 1
                attachment_count += len(exported_attachments)
                manifest.write(
                    json.dumps(
                        {
                            "form_id": form["form_id"],
                            "title": form["title"],
                            "category_name": form.get("category_name"),
                            "markdown_path": str(md_path.relative_to(export_root)),
                            "attachments": [
                                manifest_attachment(a, export_root) for a in exported_attachments
                            ],
                            "extraction_warning": warning,
                        },
                        ensure_ascii=False,
                    )
                    + "\n"
                )
                exported += 1
        write_export_readme(export_root, exported, attachment_count)
    finally:
        con.close()
    return {
        "status": "ok",
        "export_dir": str(export_root),
        "manifest": str(manifest_path),
        "forms": exported,
        "attachments": attachment_count,
        "downloaded_files": downloaded,
        "download_errors": download_errors,
        "forms_with_extraction_warnings": warnings,
    }


def sync_forms(args: argparse.Namespace) -> dict[str, Any]:
    con = connect(args.db)
    ensure_schema(con)
    sync_id = f"sync-{now_epoch_ms()}"
    con.execute(
        "INSERT INTO sync_runs(sync_id, source_url) VALUES (?, ?)",
        (sync_id, SOURCE_URL),
    )
    con.commit()
    category_names: dict[str, str] = {}
    total_count = 0
    synced = 0
    downloaded = 0
    download_errors: list[dict[str, str]] = []
    status = "complete"
    message = ""

    try:
        page_no = 1
        while True:
            data = fetch_page(page_no, args.page_size, args.category, args.search_word)
            categories = data.get("dlt_mcodesList") or []
            upsert_categories(con, categories)
            category_names.update(
                {
                    (c.get("scode") or "").strip(): (c.get("scodeKname") or "").strip()
                    for c in categories
                    if c.get("scode") and c.get("scodeKname")
                }
            )
            search_meta = data.get("dma_search") or {}
            total_count = int(search_meta.get("totalCnt") or total_count or 0)
            rows = data.get("dlt_nboardList") or []
            if args.limit:
                rows = rows[: max(0, args.limit - synced)]
            synced += store_rows(con, rows, category_names, page_no)
            con.commit()

            if args.download != "none":
                kinds = None if args.download == "all" else {args.download}
                attachment_rows = attachments_for_recent_page(con, page_no, kinds)
                for attachment in attachment_rows:
                    try:
                        download_attachment(con, attachment, args.files_dir, force=args.force)
                        downloaded += 1
                    except Exception as exc:
                        if not args.continue_on_error:
                            raise
                        download_errors.append(
                            {
                                "attachment_id": attachment["attachment_id"],
                                "original_name": attachment["original_name"],
                                "error": str(exc),
                            }
                        )
                    if args.delay:
                        time.sleep(args.delay)
                con.commit()

            if not rows:
                break
            if args.limit and synced >= args.limit:
                break
            if total_count and page_no * args.page_size >= total_count:
                break
            page_no += 1
            if args.delay:
                time.sleep(args.delay)

        rebuild_fts(con)
        con.execute(
            """
            UPDATE sync_runs
            SET completed_at = CURRENT_TIMESTAMP,
                total_count = ?,
                synced_forms = ?,
                downloaded_files = ?,
                status = ?
            WHERE sync_id = ?
            """,
            (total_count, synced, downloaded, status, sync_id),
        )
        con.commit()
    except Exception as exc:
        status = "failed"
        message = str(exc)
        con.execute(
            """
            UPDATE sync_runs
            SET completed_at = CURRENT_TIMESTAMP,
                total_count = ?,
                synced_forms = ?,
                downloaded_files = ?,
                status = ?,
                message = ?
            WHERE sync_id = ?
            """,
            (total_count, synced, downloaded, status, message, sync_id),
        )
        con.commit()
        raise
    finally:
        con.close()

    return {
        "sync_id": sync_id,
        "status": status,
        "source_url": SOURCE_URL,
        "copyright_policy_url": COPYRIGHT_POLICY_URL,
        "total_count": total_count,
        "synced_forms": synced,
        "downloaded_files": downloaded,
        "download_errors": download_errors,
        "db": str(args.db),
    }


def attachments_for_recent_page(
    con: sqlite3.Connection,
    page_no: int,
    kinds: set[str] | None,
) -> list[sqlite3.Row]:
    sql = """
        SELECT a.*
        FROM attachments a
        JOIN forms f ON f.form_id = a.form_id
        WHERE f.source_page = ?
    """
    params: list[Any] = [page_no]
    if kinds is not None:
        placeholders = ",".join("?" * len(kinds))
        sql += f" AND a.kind IN ({placeholders})"
        params.extend(sorted(kinds))
    return list(con.execute(sql, params))


def normalize(value: str) -> str:
    return " ".join(value.lower().split())


def token_list(query: str) -> list[str]:
    return [normalize(token) for token in TOKEN_RE.findall(query) if token.strip()]


def load_forms_for_search(con: sqlite3.Connection) -> list[dict[str, Any]]:
    rows = con.execute(
        """
        SELECT
          f.*,
          COALESCE(GROUP_CONCAT(a.original_name, ' '), '') AS attachment_names
        FROM forms f
        LEFT JOIN attachments a ON a.form_id = f.form_id
        GROUP BY f.form_id
        """
    ).fetchall()
    return [dict(r) for r in rows]


def attachments_for_form(con: sqlite3.Connection, form_id: str) -> list[dict[str, Any]]:
    rows = con.execute(
        """
        SELECT attachment_id, kind, original_name, system_name, download_url,
               local_path, size_bytes, sha256, downloaded_at
        FROM attachments
        WHERE form_id = ?
        ORDER BY
          CASE kind WHEN 'hwp' THEN 1 WHEN 'hwpx' THEN 2 WHEN 'doc' THEN 3
                    WHEN 'docx' THEN 4 WHEN 'pdf' THEN 5 ELSE 9 END,
          original_name
        """,
        (form_id,),
    ).fetchall()
    return [dict(r) for r in rows]


def score_form(row: dict[str, Any], query: str, tokens: list[str]) -> float:
    title = normalize(row.get("title") or "")
    category = normalize(row.get("category_name") or "")
    attachments = normalize(row.get("attachment_names") or "")
    haystack = f"{title} {category} {attachments}"
    q = normalize(query)
    score = 0.0
    if q and q in title:
        score += 20.0
    if q and q in attachments:
        score += 8.0
    for token in tokens:
        if token in title:
            score += 5.0
        if token in attachments:
            score += 2.0
        if token in category:
            score += 1.0
    if tokens and all(token in haystack for token in tokens):
        score += 3.0
    return score


def search_db(
    con: sqlite3.Connection,
    query: str,
    top_k: int,
    category: str = "",
    with_attachments: bool = True,
) -> list[dict[str, Any]]:
    tokens = token_list(query)
    category_norm = normalize(category)
    ranked: list[tuple[float, dict[str, Any]]] = []
    for row in load_forms_for_search(con):
        if category_norm:
            category_haystack = normalize(
                f"{row.get('category_code') or ''} {row.get('category_name') or ''}"
            )
            if category_norm not in category_haystack:
                continue
        score = score_form(row, query, tokens)
        if score <= 0:
            continue
        ranked.append((score, row))
    ranked.sort(key=lambda item: (-item[0], item[1].get("title") or ""))

    results = []
    for score, row in ranked[:top_k]:
        form = {
            "form_id": row["form_id"],
            "score": round(score, 3),
            "title": row["title"],
            "category_code": row["category_code"],
            "category_name": row["category_name"],
            "source_url": row["source_url"],
        }
        if with_attachments:
            form["attachments"] = attachments_for_form(con, row["form_id"])
        results.append(form)
    return results


def search_forms(args: argparse.Namespace) -> dict[str, Any]:
    con = connect(args.db)
    ensure_schema(con)
    try:
        results = search_db(con, args.query, args.top_k, args.category)
    finally:
        con.close()
    return {
        "query": args.query,
        "category": args.category or None,
        "count": len(results),
        "results": results,
        "source": "대한민국 법원 전자소송포털 양식모음",
        "source_url": SOURCE_URL,
    }


def select_download_attachments(
    con: sqlite3.Connection,
    form_id: str,
    kind: str,
) -> list[sqlite3.Row]:
    if kind == "all":
        return list(con.execute("SELECT * FROM attachments WHERE form_id = ?", (form_id,)))
    return list(
        con.execute(
            "SELECT * FROM attachments WHERE form_id = ? AND kind = ?",
            (form_id, kind),
        )
    )


def download_forms(args: argparse.Namespace) -> dict[str, Any]:
    con = connect(args.db)
    ensure_schema(con)
    try:
        form_id = args.form_id
        if not form_id:
            matches = search_db(con, args.query, 1, args.category, with_attachments=False)
            if not matches:
                raise CourtFormsError(f"no court form matched query: {args.query}")
            form_id = matches[0]["form_id"]
        form = con.execute("SELECT * FROM forms WHERE form_id = ?", (form_id,)).fetchone()
        if not form:
            raise CourtFormsError(f"form_id not found: {form_id}")
        attachment_rows = select_download_attachments(con, form_id, args.kind)
        if not attachment_rows:
            raise CourtFormsError(f"no {args.kind} attachment found for {form_id}")
        downloaded = [
            download_attachment(con, row, args.out_dir, force=args.force)
            for row in attachment_rows
        ]
        con.commit()
    finally:
        con.close()
    return {
        "form_id": form_id,
        "title": form["title"],
        "source": "대한민국 법원 전자소송포털 양식모음",
        "source_url": SOURCE_URL,
        "copyright_policy_url": COPYRIGHT_POLICY_URL,
        "downloaded": downloaded,
    }


def info(args: argparse.Namespace) -> dict[str, Any]:
    con = connect(args.db)
    ensure_schema(con)
    try:
        forms = con.execute("SELECT COUNT(*) FROM forms").fetchone()[0]
        attachments = con.execute("SELECT COUNT(*) FROM attachments").fetchone()[0]
        downloaded = con.execute(
            "SELECT COUNT(*) FROM attachments WHERE local_path IS NOT NULL"
        ).fetchone()[0]
        last_sync = con.execute(
            """
            SELECT sync_id, started_at, completed_at, total_count, synced_forms,
                   downloaded_files, status, message
            FROM sync_runs
            ORDER BY started_at DESC
            LIMIT 1
            """
        ).fetchone()
    finally:
        con.close()
    return {
        "status": "ok",
        "db": str(args.db),
        "files_dir": str(args.files_dir),
        "forms": forms,
        "attachments": attachments,
        "downloaded_files": downloaded,
        "last_sync": dict(last_sync) if last_sync else None,
        "source_url": SOURCE_URL,
        "copyright_policy_url": COPYRIGHT_POLICY_URL,
    }


def add_common_args(parser: argparse.ArgumentParser) -> None:
    parser.add_argument("--db", type=Path, default=DB_PATH)
    parser.add_argument("--files-dir", type=Path, default=FILES_DIR)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Korean court form DB helper.")
    sub = parser.add_subparsers(dest="command", required=True)

    init_p = sub.add_parser("init", help="Create or upgrade the SQLite schema.")
    add_common_args(init_p)

    sync_p = sub.add_parser("sync", help="Sync public form metadata from ecfs.scourt.go.kr.")
    add_common_args(sync_p)
    sync_p.add_argument("--page-size", type=int, default=30)
    sync_p.add_argument("--limit", type=int, default=0)
    sync_p.add_argument("--category", default="", help="Category code, e.g. a=민사, i=형사.")
    sync_p.add_argument("--search-word", default="")
    sync_p.add_argument("--download", choices=["none", "hwp", "hwpx", "doc", "docx", "pdf", "all"], default="none")
    sync_p.add_argument("--delay", type=float, default=0.25)
    sync_p.add_argument("--force", action="store_true")
    sync_p.add_argument("--continue-on-error", action="store_true")

    search_p = sub.add_parser("search", help="Search synced form metadata.")
    add_common_args(search_p)
    search_p.add_argument("query")
    search_p.add_argument("--top-k", type=int, default=5)
    search_p.add_argument("--category", default="")

    download_p = sub.add_parser("download", help="Download template files for one form.")
    add_common_args(download_p)
    download_p.add_argument("--form-id", default="")
    download_p.add_argument("--query", default="")
    download_p.add_argument("--category", default="")
    download_p.add_argument("--kind", choices=["hwp", "hwpx", "doc", "docx", "pdf", "all"], default="hwp")
    download_p.add_argument("--out-dir", type=Path, default=FILES_DIR)
    download_p.add_argument("--force", action="store_true")

    export_p = sub.add_parser(
        "export-md",
        help="Export synced forms as repo-friendly Markdown folders.",
    )
    add_common_args(export_p)
    export_p.add_argument("--output", type=Path, default=Path("data/court-forms"))
    export_p.add_argument("--copy-files", action="store_true")
    export_p.add_argument("--download-missing", action="store_true")
    export_p.add_argument("--delay", type=float, default=0.25)
    export_p.add_argument("--force", action="store_true")
    export_p.add_argument("--continue-on-error", action="store_true")

    info_p = sub.add_parser("info", help="Show local DB status.")
    add_common_args(info_p)
    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)
    try:
        if args.command == "init":
            con = connect(args.db)
            ensure_schema(con)
            con.close()
            result = {"status": "ok", "db": str(args.db)}
        elif args.command == "sync":
            if args.page_size < 1 or args.page_size > 100:
                raise CourtFormsError("--page-size must be between 1 and 100")
            result = sync_forms(args)
        elif args.command == "search":
            result = search_forms(args)
        elif args.command == "download":
            if not args.form_id and not args.query:
                raise CourtFormsError("download requires --form-id or --query")
            result = download_forms(args)
        elif args.command == "export-md":
            result = export_markdown(args)
        elif args.command == "info":
            result = info(args)
        else:
            parser.error("unknown command")
            return 2
    except CourtFormsError as exc:
        print(f"court-forms: {exc}", file=sys.stderr)
        return 1
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
